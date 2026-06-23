"""P3.3 — Lecteur slips / memos courtiers (PDF + Word).

Extrait tables et texte de :
- PDF slips (Aon, GuyCarp, Willis Re, Howden, Lockton)
- Word memos (renouvellement, conditions, structures)

Détecte automatiquement :
- Company / Broker / Inception / Period / Currency
- Treaty structure : Share / Limit / Deductible / Reinstatement / Commission
- Wording clauses présents

Limite : PDF scannés (image only) nécessitent OCR (pas implémenté ici).
"""
from __future__ import annotations
from dataclasses import dataclass, field
from pathlib import Path
import re
import pandas as pd


# ─── Patterns de détection ────────────────────────────────────────────────────

FIELD_PATTERNS = {
    "company":      [r"(?:Reinsured|Insured|Company|Cedant|Cédante)\s*[:\-]\s*(.+?)(?:\n|$)"],
    "broker":       [r"(?:Broker|Courtier|Intermediary)\s*[:\-]\s*(.+?)(?:\n|$)"],
    "leader":       [r"(?:Leader|Lead Reinsurer|Apériteur)\s*[:\-]\s*(.+?)(?:\n|$)"],
    "inception":    [r"(?:Inception|Effective Date|Date d['']effet)\s*[:\-]\s*(.+?)(?:\n|$)"],
    "expiry":       [r"(?:Expiry|Expiration|Date de fin)\s*[:\-]\s*(.+?)(?:\n|$)"],
    "currency":     [r"(?:Currency|Devise|Cur\.)\s*[:\-]\s*([A-Z]{3})"],
    "period_basis": [r"\b(RAD|LOD|Risks?\s+Attaching|Losses?\s+Occurring)\b"],
    "country":      [r"(?:Country|Territory|Pays)\s*[:\-]\s*(.+?)(?:\n|$)"],
    "share":        [r"(?:Share|Part)\s*[:\-]?\s*(\d+(?:\.\d+)?)\s*%"],
    "treaty_type":  [r"\b(Quote\s+Share|QS|Surplus|Excess of Loss|XL|XOL|Stop\s+Loss|Facility)\b"],
    "limit":        [r"(?:Limit|Capacity|Garantie)\s*[:\-]?\s*(?:USD|EUR|EGP|GBP|MAD)?\s*([\d,\.\s]+)"],
    "deductible":   [r"(?:Deductible|Priority|Retention|Franchise)\s*[:\-]?\s*(?:USD|EUR)?\s*([\d,\.\s]+)"],
    "reinstatement":[r"(\d+)\s*(?:@|at)\s*(\d+(?:\.\d+)?)\s*%"],
    "commission":   [r"(?:Commission|COM)\s*[:\-]?\s*(\d+(?:\.\d+)?)\s*%"],
    "brokerage":    [r"(?:Brokerage|Courtage)\s*[:\-]?\s*(\d+(?:\.\d+)?)\s*%"],
}

WORDING_KEYWORDS = [
    "NMA", "LMA", "Hours Clause", "Sanctions", "Terrorism",
    "Cyber", "Nuclear", "War", "Arbitration", "Currency Clause",
    "Follow the Fortunes", "Errors and Omissions", "Ultimate Net Loss",
    "Reinstatement", "Communicable Disease", "Pollution",
]


@dataclass
class ExtractedDocument:
    """Résultat extraction d'un slip / memo."""
    source_path: str
    file_type: str               # 'pdf' / 'docx'
    raw_text: str = ""
    n_pages: int = 0
    tables: list = field(default_factory=list)   # liste de DataFrame
    detected_fields: dict = field(default_factory=dict)
    detected_wording: list = field(default_factory=list)
    warnings: list = field(default_factory=list)


def read_pdf(path: str | Path) -> ExtractedDocument:
    """Lit un PDF avec pdfplumber : texte + tables."""
    import pdfplumber
    path = Path(path)
    doc = ExtractedDocument(source_path=str(path), file_type='pdf')

    try:
        with pdfplumber.open(path) as pdf:
            doc.n_pages = len(pdf.pages)
            all_text = []
            for page in pdf.pages:
                t = page.extract_text() or ""
                all_text.append(t)
                # Extract tables
                for tbl in page.extract_tables():
                    if tbl and len(tbl) > 1:
                        df = pd.DataFrame(tbl[1:], columns=tbl[0])
                        doc.tables.append(df)
            doc.raw_text = "\n".join(all_text)
    except Exception as e:
        doc.warnings.append(f"PDF parse error: {e}")
        return doc

    if not doc.raw_text.strip():
        doc.warnings.append("PDF semble être scanné (aucun texte). OCR requis.")
        return doc

    doc.detected_fields = _extract_fields(doc.raw_text)
    doc.detected_wording = _detect_wording(doc.raw_text)
    return doc


def read_docx(path: str | Path) -> ExtractedDocument:
    """Lit un Word avec python-docx : paragraphes + tables."""
    from docx import Document
    path = Path(path)
    doc = ExtractedDocument(source_path=str(path), file_type='docx')

    try:
        d = Document(path)
        paragraphs = [p.text for p in d.paragraphs if p.text.strip()]
        doc.raw_text = "\n".join(paragraphs)
        doc.n_pages = max(1, len(paragraphs) // 30)  # heuristic
        for tbl in d.tables:
            rows = []
            for r in tbl.rows:
                rows.append([c.text.strip() for c in r.cells])
            if rows and len(rows) > 1:
                df = pd.DataFrame(rows[1:], columns=rows[0])
                doc.tables.append(df)
    except Exception as e:
        doc.warnings.append(f"DOCX parse error: {e}")
        return doc

    doc.detected_fields = _extract_fields(doc.raw_text)
    doc.detected_wording = _detect_wording(doc.raw_text)
    return doc


def _extract_fields(text: str) -> dict:
    """Applique tous les patterns sur le texte et retourne champs détectés."""
    out = {}
    for field, patterns in FIELD_PATTERNS.items():
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE | re.MULTILINE)
            if m:
                if field == "reinstatement" and len(m.groups()) >= 2:
                    out[field] = f"{m.group(1)}@{m.group(2)}%"
                else:
                    out[field] = m.group(1).strip() if m.groups() else m.group(0).strip()
                break
    return out


def _detect_wording(text: str) -> list:
    """Liste les clauses de wording trouvées dans le texte."""
    found = []
    upper = text.upper()
    for kw in WORDING_KEYWORDS:
        if kw.upper() in upper:
            # Extraire ligne contenant le keyword
            for line in text.split("\n"):
                if kw.upper() in line.upper():
                    found.append({"keyword": kw, "context": line.strip()[:200]})
                    break
    return found


def read_document(path: str | Path) -> ExtractedDocument:
    """Dispatcher : lit PDF ou DOCX selon extension."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == '.pdf':
        return read_pdf(path)
    elif suffix in ('.docx', '.doc'):
        return read_docx(path)
    else:
        raise ValueError(f"Format non supporté pour document: {suffix}")


def merge_into_uw_header(doc: ExtractedDocument, header) -> None:
    """Injecte les champs détectés dans un TreatyHeader existant.

    Préserve les valeurs déjà saisies par l'utilisateur si présentes.
    """
    fields = doc.detected_fields
    if not header.company and fields.get("company"):
        header.company = fields["company"]
    if not header.broker and fields.get("broker"):
        header.broker = fields["broker"]
    if fields.get("leader") and not header.leader.startswith("SMGA"):
        header.leader = fields["leader"]
    if fields.get("inception"):
        header.inception_date = fields["inception"]
    if fields.get("expiry"):
        header.expiry_date = fields["expiry"]
    if fields.get("currency"):
        header.original_currency = fields["currency"]
    if fields.get("country"):
        header.country = fields["country"]


def find_treaty_tables(doc: ExtractedDocument) -> list:
    """Filtre les tables extraites pour ne garder que celles ressemblant à
    des tableaux de structure traité.

    Heuristique : header contient au moins 2 mots-clés parmi
    [Treaty, Share, Limit, EPI, Deductible, Premium, LR].
    """
    treaty_keywords = ['treaty', 'share', 'limit', 'epi', 'gnpi',
                        'deductible', 'priority', 'premium', 'lr',
                        'loss ratio', 'commission', 'brokerage',
                        'reinstatement', 'layer']
    out = []
    for df in doc.tables:
        cols_str = ' '.join(str(c).lower() for c in df.columns)
        n_match = sum(1 for kw in treaty_keywords if kw in cols_str)
        if n_match >= 2:
            out.append(df)
    return out


def summarize_document(doc: ExtractedDocument) -> dict:
    """Résumé pour affichage Streamlit."""
    treaty_tables = find_treaty_tables(doc)
    return {
        "file": Path(doc.source_path).name,
        "type": doc.file_type,
        "pages": doc.n_pages,
        "text_length_chars": len(doc.raw_text),
        "n_tables_extracted": len(doc.tables),
        "n_treaty_tables_detected": len(treaty_tables),
        "fields_detected": doc.detected_fields,
        "n_wording_clauses": len(doc.detected_wording),
        "wording_clauses": [w["keyword"] for w in doc.detected_wording],
        "warnings": doc.warnings,
        "treaty_tables_preview": [df.head().to_dict('records') for df in treaty_tables[:3]],
    }
